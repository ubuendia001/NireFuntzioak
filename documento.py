import pandas as pd
import numpy as np
import docx
from docx import Document
from docx.shared import Mm
from docxtpl import DocxTemplate
import os

def mandar_a_impresora(archivo_nombre):
    """
    Esta función imprime un archivo en la impresora que este definida como predefinida.
    Importaciones necesarias:
        -import os.
    Parametros:
        -archivo_nombre(string): Nombre (con extensión) del archivo que se quiere imprimir.

    """
    os.startfile(archivo_nombre, "print")


def tabla_a_documento(tabla, documento,incluir_index=True):
    """
    Esta función pega en el documento word indicado el DataFrame indicado como tabla.
    Importaciones necesarias:
        -import docx.
    Parametros:
        -tabla(pandas.DataFrame): tabla del tipo DataFrame que se quiere pegar en el word
        -documento(docx.Document()): Documento word en el que se quiere pegar la tabla.

    """
    if incluir_index==True:
        filas = tabla.shape[0]
        columnas = tabla.shape[1]
        tabla_trabajo = documento.add_table(rows=filas+1, cols=columnas+1)
        fila = 0
        while fila < filas:
            tabla_trabajo.cell(fila+1, 0).text = str(tabla.index[fila])
            fila += 1

        columna = 0
        while columna < columnas:
            tabla_trabajo.cell(0, columna+1).text = str(tabla.columns[columna])
            columna += 1

        fila = 1
        while fila <= filas:
            columna = 1
            while columna <= columnas:
                tabla_trabajo.cell(fila, columna).text = str(
                    tabla.iloc[fila-1, columna-1])
                columna += 1
            fila += 1

        tabla_trabajo.style = 'Table Grid'
    else:
        filas = tabla.shape[0]
        columnas = tabla.shape[1]
        tabla_trabajo = documento.add_table(rows=filas+1, cols=columnas)

        columna = 0
        while columna < columnas:
            tabla_trabajo.cell(0, columna).text = str(tabla.columns[columna])
            columna += 1

        fila = 1
        while fila <= filas:
            columna = 0
            while columna <= columnas-1:
                tabla_trabajo.cell(fila, columna).text = str(
                    tabla.iloc[fila-1, columna])
                columna += 1
            fila += 1

        tabla_trabajo.style = 'Table Grid'

def tabla_a_imagen(datos,archivo_nombre,incluir_index=True):
    """
    Esta función pasa un DataFrame a una imagen del formato .png.
    Parametros:
        -datos(pandas.DataFrame): DataFrame de pandas que se quiere pasar a la imagen png.
        -archivo_nombre(string): Nombre (sin extensión) para la imagen creada.

    """
    tabla_datos=datos.to_numpy()
    tabla_cabeceras=datos.columns
    tabla_indices=datos.index

    fig,ax=plt.subplots(figsize=(5,5))
    ax.get_xaxis().set_visible(False)
    ax.get_yaxis().set_visible(False)
    plt.box(on=None)

    colores_cabecera=['skyblue']*len(tabla_cabeceras)
    colores_filas=['skyblue']*len(tabla_indices)

    if incluir_index==False:
        rowLabels=['']*len(tabla_indices)

    tabla=plt.table(loc='center',cellText=tabla_datos,colLabels=tabla_cabeceras,rowLabels=tabla_indices,
                colColours=colores_cabecera,rowColours=colores_filas)

    tabla.set_fontsize(14)
    plt.tight_layout()
    plt.savefig(archivo_nombre+'.png',dpi=300)

#No sé si esto funciona cuando las imagenes falsas están en otra carpeta
def generar_informe(rute,variables_dict,imagenes_dict,nombre_informe):
    """
    Esta función genera un informe a partir de una plantilla ubicada en el directorio "rute"
    actualizando los datos con variables_dict y los gráficos con imagenes_dict de la plantilla
    y guardándolo con el nombre nombre_informe
    Importaciones necesarias:
        -import pandas as pd
        -import numpy as np
        -from docx import Document
        -from docx.shared import Mm
        -from docxtpl import DocxTemplate
        -import os

    Parametros:
        -rute(string): ruta donde esta la plantilla word
        -variables_dict(string): Diccionario donde se indica el listado de variables y sus valores.
        -imagenes_dict(string): Diccionario donde se indica qué imagen va en cada sitio.
        -nombre_informe(string):Nombre con el que queremos guardar el informe (sin la extensión)
    """
    documento=DocxTemplate(rute + "Plantilla.docx")
    documento.render(variables_dict)
    for foto_antigua,foto_nueva in imagenes_dict.items():
        documento.replace_pic(foto_antigua,foto_nueva)
    documento.save(rute + nombre_informe+".docx")

#/////////////////////////////////////////////////////////////////////////////////////////////////
# --------------  CREAR UN ARCHIVO WORD, AÑADIr DIFERENTES ELEMENTOS Y CONVERTIRLO A PDF------------------------------------

"""
import docx
from docx.shared import Mm

documento = docx.Document() #Crear el documento
documento.add_heading(texto, level=0) # Meter título de nivel 0 (el mayor) en el documento
document.add_paragraph(texto, style='List Bullet') # Meter párrafo en el documento como lista de puntos
document.add_paragraph(texto,style='List Number') # Meter párrafo en el documento como lista numerada
document.add_paragraph(texto,style='Intense Quote') # Meter párrafo en el documento casi como título
parrafo = document.add_paragraph(texto) # Meter párrafo en el documento
parrafo.add_run(texto).bold = True # Añadir frase en negrita al párrafo anteriormente introducido
parrafo.add_run(texto) # Añadir frase al párrafo anteriormente introducido
parrafo.add_run(texto).italic = True# Añadir frase en cursiva al párrafo anteriormente introducido
tabla_a_documento(nombre_tabla, documento) # Meter dataframe de pandas de nombre 'nombre_tabla' en el documento
documento.add_page_break() #Para pasar a la siguiente hoja
documento.add_picture(directorio_imagen,width=Mm(120),height=Mm(120)) # Meter  indicando el directorio y el ancho o el alto (o los dos)
documento.save(directorio_documento) # Guardar el documento
"""


"""
instalación: pip install docx2pdf
from docx2pdf import convert
convert('C:/Users/urkob/Desktop/Prueba')
Se puede indicar un carpeta para que convierta todos los documentos, o un documento específicamente.
"""
